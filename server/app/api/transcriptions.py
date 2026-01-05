from fastapi import APIRouter, Depends, HTTPException, Query, BackgroundTasks, status
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import or_
from typing import List, Optional
from pathlib import Path
from uuid import UUID
import secrets
import base64
from app.api.deps import get_db, get_current_db_user, require_active, require_admin
from app.core.supabase import get_current_active_user
from app.core.config import settings
from app.models.transcription import Transcription
from app.models.user import User
from app.models.summary import Summary
from app.models.chat_message import ChatMessage
from app.models.share_link import ShareLink
from app.models.channel import Channel, ChannelMembership, TranscriptionChannel
from app.schemas.transcription import Transcription as TranscriptionSchema, PaginatedResponse
from app.schemas.summary import Summary as SummarySchema
from app.schemas.chat import ChatMessage as ChatMessageSchema, ChatHistoryResponse
from app.schemas.share import ShareLink as ShareLinkSchema
from app.schemas.admin import TranscriptionChannelAssignmentRequest, ChannelResponse

import logging
import asyncio
import time

logger = logging.getLogger(__name__)

router = APIRouter()


# ============================================================================
# Helper Functions
# ============================================================================

def _format_fake_srt(text: str) -> str:
  """
  Generate fake SRT format from plain text (fallback for old transcriptions).

  Splits text by newlines and assigns sequential fake timestamps.

  Args:
    text: Plain transcription text

  Returns:
    SRT-formatted string with fake timestamps
  """
  lines = text.strip().split("\n")
  srt_lines = []
  for i, line in enumerate(lines, 1):
    if line.strip():
      # Fake timestamp: each line gets 1 second
      start_time = f"{i:02d}:00:00,000"
      end_time = f"{i:02d}:00:01,000"
      srt_lines.append(f"{i}\n{start_time} --> {end_time}\n{line}\n")
  return "\n".join(srt_lines)


# ============================================================================
# API Endpoints
# ============================================================================

@router.get("", response_model=PaginatedResponse[TranscriptionSchema])
async def list_transcriptions(
    page: int = Query(1, ge=1, description="Page number (1-indexed)"),
    page_size: int = Query(None, ge=1, le=settings.MAX_PAGE_SIZE, description="Number of items per page"),
    stage: str = Query(None, description="Filter by stage: uploading, transcribing, summarizing, completed, failed"),
    channel_id: str = Query(None, description="Filter by channel ID (regular users only)"),
    db: Session = Depends(get_db),
    current_db_user: User = Depends(get_current_db_user)
):
    """
    文字起こしリストを取得

    - Regular users: own content + content from assigned channels
    - Admin users: all content (bypasses channel filters)
    """
    # Use default page size from config if not specified
    if page_size is None:
        page_size = settings.DEFAULT_PAGE_SIZE

    # Build base query
    # Admin sees everything, regular users see own + channel-assigned
    if current_db_user.is_admin:
        query = db.query(Transcription)
    else:
        # Get user's channel IDs
        user_channel_ids = db.query(ChannelMembership.channel_id).filter(
            ChannelMembership.user_id == current_db_user.id
        ).all()
        user_channel_ids = [c[0] for c in user_channel_ids]

        # Get transcription IDs from user's channels
        channel_transcription_ids = db.query(TranscriptionChannel.transcription_id).filter(
            TranscriptionChannel.channel_id.in_(user_channel_ids)
        ).all()
        channel_transcription_ids = [t[0] for t in channel_transcription_ids]

        # Query: own OR in channels
        query = db.query(Transcription).filter(
            or_(
                Transcription.user_id == current_db_user.id,
                Transcription.id.in_(channel_transcription_ids)
            )
        )

        # Optional channel filter (for regular users)
        if channel_id:
            # Verify user is member of this channel
            is_member = db.query(ChannelMembership).filter(
                ChannelMembership.channel_id == channel_id,
                ChannelMembership.user_id == current_db_user.id
            ).first()
            if not is_member:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail="Not a member of this channel"
                )
            # Filter to only show content from this channel
            query = query.join(TranscriptionChannel).filter(
                TranscriptionChannel.channel_id == channel_id
            )

    if stage:
        query = query.filter(Transcription.stage == stage)

    # Get total count
    total = query.count()

    # Calculate offset and apply pagination
    offset = (page - 1) * page_size
    transcriptions = query.order_by(Transcription.created_at.desc()).offset(offset).limit(page_size).all()

    # Calculate total pages
    total_pages = (total + page_size - 1) // page_size if total > 0 else 0

    return PaginatedResponse(
        total=total,
        page=page,
        page_size=page_size,
        total_pages=total_pages,
        data=transcriptions
    )


@router.get("/{transcription_id}", response_model=TranscriptionSchema)
async def get_transcription(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    文字起こし詳細を取得
    """
    # Convert string ID to UUID for database query
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()
    
    if not transcription:
        raise HTTPException(status_code=404, detail="文字起こしが見つかりません")
    
    return transcription


@router.delete("/all", status_code=200)
async def delete_all_transcriptions(
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    すべての文字起こしを削除 (ファイルも含む)
    処理中の場合はキャンセルしてプロセスを終了してから削除します
    """
    from app.services.transcription_processor import (
        is_transcription_active,
        mark_transcription_cancelled,
        kill_transcription_processes,
        get_transcription_task_info
    )

    # Get all user's transcriptions
    transcriptions = db.query(Transcription).filter(
        Transcription.user_id == current_user.get("id")
    ).all()

    if not transcriptions:
        return {"deleted_count": 0, "message": "削除する項目がありません"}

    deleted_count = 0
    output_dir = Path("/app/data/output")

    for transcription in transcriptions:
        transcription_id = str(transcription.id)

        # Cancel active transcriptions
        if is_transcription_active(transcription_id):
            logger.info(f"[DELETE ALL] Cancelling active transcription: {transcription_id}")
            mark_transcription_cancelled(transcription_id)
            kill_transcription_processes(transcription_id)
            time.sleep(0.1)  # Brief pause between cancellations

        # Delete files
        try:
            import os
            from app.services.storage_service import get_storage_service

            # Delete from Supabase Storage
            if transcription.storage_path:
                try:
                    storage_service = get_storage_service()
                    storage_service.delete_transcription_text(str(transcription.id))
                    storage_service.delete_transcription_segments(str(transcription.id))
                    storage_service.delete_original_output(str(transcription.id))
                    storage_service.delete_formatted_text(str(transcription.id))
                    storage_service.delete_notebooklm_guideline(str(transcription.id))
                    logger.info(f"[DELETE ALL] Deleted from storage: {transcription.storage_path}")
                except Exception as e:
                    logger.warning(f"[DELETE ALL] Failed to delete from storage: {e}")

            if transcription.file_path and os.path.exists(transcription.file_path):
                os.remove(transcription.file_path)

            for ext in [".wav", ".txt", ".srt", ".vtt", ".json"]:
                output_file = output_dir / f"{transcription.id}{ext}"
                if output_file.exists():
                    output_file.unlink()

                converted_wav = output_dir / f"{transcription.id}_converted.wav"
                if converted_wav.exists():
                    converted_wav.unlink()
        except Exception as e:
            logger.error(f"[DELETE ALL] File deletion error for {transcription_id}: {e}")

        db.delete(transcription)
        deleted_count += 1

    db.commit()
    logger.info(f"[DELETE ALL] Deleted {deleted_count} transcriptions for user {current_user.get('id')}")

    return {
        "deleted_count": deleted_count,
        "message": f"{deleted_count}件の転写を削除しました"
    }


@router.delete("/{transcription_id}", status_code=204)
async def delete_transcription(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    文字起こしを削除 (ファイルも含む)
    処理中の場合はキャンセルしてプロセスを終了してから削除します
    """
    # Convert string ID to UUID for database query
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="文字起こしが見つかりません")

    # Check if transcription is currently processing and cancel if needed
    from app.services.transcription_processor import (
        is_transcription_active,
        mark_transcription_cancelled,
        kill_transcription_processes,
        get_transcription_task_info
    )

    if is_transcription_active(transcription_id):
        logger.info(f"[DELETE] Cancelling active transcription: {transcription_id}")

        # Get task info before cancelling
        task_info = get_transcription_task_info(transcription_id)
        if task_info:
            pids = task_info.get("pids", set())
            logger.info(f"[DELETE] Task info for {transcription_id}: Stage={task_info.get('stage')}, PIDs={list(pids)}")

        # Mark as cancelled (signals the transcription thread)
        was_marked = mark_transcription_cancelled(transcription_id)
        if was_marked:
            logger.info(f"[DELETE] Marked transcription {transcription_id} as cancelled")

        # Kill subprocesses (FFmpeg, whisper-cli)
        killed_count = kill_transcription_processes(transcription_id)
        logger.info(f"[DELETE] Killed {killed_count} subprocesses for transcription {transcription_id}")

        # Wait a moment for processes to terminate
        time.sleep(0.5)

    # ファイル削除
    try:
        import os
        from pathlib import Path
        from app.services.storage_service import get_storage_service

        # Delete from Supabase Storage
        if transcription.storage_path:
            try:
                storage_service = get_storage_service()
                storage_service.delete_transcription_text(str(transcription.id))
                storage_service.delete_transcription_segments(str(transcription.id))
                storage_service.delete_original_output(str(transcription.id))
                storage_service.delete_formatted_text(str(transcription.id))
                storage_service.delete_notebooklm_guideline(str(transcription.id))
                logger.info(f"[DELETE] Deleted from storage: {transcription.storage_path}")
            except Exception as e:
                logger.warning(f"[DELETE] Failed to delete from storage: {e}")

        # アップロードファイル削除
        if transcription.file_path and os.path.exists(transcription.file_path):
            os.remove(transcription.file_path)
            logger.info(f"[DELETE] Deleted upload file: {transcription.file_path}")

        # 出力ファイル削除 (wav, txt, srt, vtt, json)
        output_dir = Path("/app/data/output")
        for ext in [".wav", ".txt", ".srt", ".vtt", ".json"]:
            output_file = output_dir / f"{transcription.id}{ext}"
            if output_file.exists():
                output_file.unlink()
                logger.info(f"[DELETE] Deleted output file: {output_file}")

            # Also check for converted wav
            converted_wav = output_dir / f"{transcription.id}_converted.wav"
            if converted_wav.exists():
                converted_wav.unlink()
                logger.info(f"[DELETE] Deleted converted wav: {converted_wav}")

    except Exception as e:
        logger.error(f"ファイル削除エラー: {e}")
        # DB削除は続行する
        pass

    db.delete(transcription)
    db.commit()
    logger.info(f"[DELETE] Deleted transcription {transcription_id} from database")
    return None


@router.get("/{transcription_id}/download")
async def download_transcription(
  transcription_id: str,
  format: str = Query("txt", pattern="^(txt|srt|formatted)$"),
  db: Session = Depends(get_db),
  current_user: dict = Depends(get_current_active_user)
):
  """
  下载转录文件

  文件按需生成:
  - txt: 原始转录文本
  - formatted: LLM格式化后的文本（带标点符号和段落）
  - srt: 从存储的转录文本生成（带时间戳）

  Args:
    transcription_id: 转录ID
    format: 文件格式 (txt, formatted, srt)

  Returns:
    StreamingResponse: 下载文件
  """
  # 认证确认 - Convert string ID to UUID for database query
  try:
    transcription_uuid = UUID(transcription_id)
  except ValueError:
    raise HTTPException(status_code=422, detail="Invalid transcription ID format")
  transcription = db.query(Transcription).filter(
    Transcription.id == transcription_uuid,
    Transcription.user_id == current_user.get("id")
  ).first()

  if not transcription:
    raise HTTPException(status_code=404, detail="未找到转录")

  # Formatted text - LLM formatted with punctuation and paragraphs
  if format == "formatted":
    from app.services.storage_service import get_storage_service
    storage_service = get_storage_service()

    if storage_service.formatted_text_exists(transcription_id):
      content = storage_service.get_formatted_text(transcription_id)
      logger.info(f"Retrieved formatted text for: {transcription_id}")
    else:
      # Formatted text not available, return original text
      content = transcription.text
      logger.info(f"Formatted text not available, returning original for: {transcription_id}")

    download_filename = f"{original_filename}_formatted.txt"

    from fastapi.responses import StreamingResponse
    from io import StringIO

    buffer = StringIO(content)
    return StreamingResponse(
      buffer,
      media_type="text/plain; charset=utf-8",
      headers={
        "Content-Disposition": f'attachment; filename="{download_filename}"'
      }
    )

  # TXT/SRT - 从存储的转录文本按需生成
  if not transcription.text:
    raise HTTPException(status_code=400, detail="转录内容为空")

  content = transcription.text
  original_filename = Path(transcription.file_name).stem

  # SRT 格式 - 使用实际时间戳（如果可用）
  if format == "srt":
    from app.services.storage_service import get_storage_service
    storage_service = get_storage_service()

    # Check if segments file exists (new transcriptions)
    if storage_service.segments_exist(transcription_id):
      segments = storage_service.get_transcription_segments(transcription_id)
      if segments:
        # Use real timestamps from segments
        srt_lines = []
        for i, segment in enumerate(segments, 1):
          srt_lines.append(f"{i}")
          srt_lines.append(f"{segment['start']} --> {segment['end']}")
          srt_lines.append(segment['text'])
          srt_lines.append("")  # Empty line between entries
        content = "\n".join(srt_lines)
        logger.info(f"Generated SRT with {len(segments)} segments from stored data")
      else:
        # Fallback to fake timestamps
        content = _format_fake_srt(content)
        logger.warning("Segments file was empty, using fake timestamps")
    else:
      # Old transcription without segments - use fake timestamps
      content = _format_fake_srt(content)
      logger.info(f"No segments file, using fake timestamps for backward compatibility")

    download_filename = f"{original_filename}.srt"
  else:
    download_filename = f"{original_filename}.txt"

  # 使用 StreamingResponse 返回内存内容
  from fastapi.responses import StreamingResponse
  from io import StringIO

  buffer = StringIO(content)

  return StreamingResponse(
    buffer,
    media_type="text/plain; charset=utf-8",
    headers={
      "Content-Disposition": f'attachment; filename="{download_filename}"'
    }
  )


@router.get("/{transcription_id}/download-docx")
async def download_summary_docx(
    transcription_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    下载AI摘要的DOCX文件

    使用python-docx将摘要Markdown转换为DOCX格式，支持中文。

    Args:
        transcription_id: 转录ID

    Returns:
        FileResponse: DOCX文件下载
    """
    # 验证转录所有权 - Convert string ID to UUID for database query
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="未找到转录")

    # 获取摘要
    if not transcription.summaries or len(transcription.summaries) == 0:
        raise HTTPException(status_code=404, detail="未找到摘要数据")

    summary_text = transcription.summaries[0].summary_text
    if not summary_text:
        raise HTTPException(status_code=400, detail="摘要内容为空")

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import tempfile
        import re

        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        docx_path = tempfile.mktemp(suffix=".docx", dir=temp_dir)

        # 创建Document对象
        doc = Document()

        # 设置默认字体（支持中文）
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        doc.styles['Normal'].font.size = Pt(11)

        # 简单的Markdown解析
        lines = summary_text.split('\n')
        in_list = False

        for line in lines:
            line = line.rstrip()

            # 空行
            if not line:
                if in_list:
                    in_list = False
                doc.add_paragraph()
                continue

            # 标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 6:
                    heading_text = line.lstrip('#').strip()
                    heading = doc.add_heading(heading_text, level=min(level, 9))
                    # 设置中文字体
                    for run in heading.runs:
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    continue

            # 列表项
            if line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\. ', line):
                if not in_list:
                    in_list = True
                list_text = line.lstrip('-*+').lstrip('0123456789.')
                p = doc.add_paragraph(list_text, style='List Bullet')
                # 设置中文字体
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                continue

            # 代码块（简单处理）
            if line.startswith('```'):
                continue

            # 普通段落 - 处理内联格式
            in_list = False
            p = doc.add_paragraph()

            # 处理粗体和斜体
            parts = re.split(r'(\*\*.*?\*\*|_.*?_)', line)
            for part in parts:
                if not part:
                    continue

                run = p.add_run(part)

                # 设置中文字体
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

                if part.startswith('**') and part.endswith('**'):
                    run.bold = True
                    run.text = part[2:-2]
                elif part.startswith('_') and part.endswith('_'):
                    run.italic = True
                    run.text = part[1:-1]

        # 保存文档
        doc.save(docx_path)

        # 设置文件名
        original_filename = Path(transcription.file_name).stem
        download_filename = f"{original_filename}-摘要.docx"

        # 定义清理函数（在响应发送后执行）
        def cleanup_temp_files():
            import shutil
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
                logger.debug(f"Cleaned up temp directory: {temp_dir}")
            except Exception as e:
                logger.warning(f"Failed to cleanup temp directory: {e}")

        # 添加后台任务在响应完成后清理
        background_tasks.add_task(cleanup_temp_files)

        return FileResponse(
            path=docx_path,
            filename=download_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx未安装，请联系管理员"
        )
    except Exception as e:
        logger.error(f"DOCX生成失败 {transcription_id}: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX生成失败: {str(e)}")


@router.get("/{transcription_id}/download-notebooklm")
async def download_notebooklm_guideline(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    下载NotebookLM演示文稿指南

    返回基于转录内容生成的NotebookLM指南文本文件。
    该指南可用于NotebookLM生成演示文稿幻灯片。

    Args:
        transcription_id: 转录ID

    Returns:
        StreamingResponse: 下载文件

    Raises:
        HTTPException: 404 - 未找到指南文件
    """
    # 验证转录所有权
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="未找到转录")

    # Check if guideline exists
    from app.services.storage_service import get_storage_service
    storage_service = get_storage_service()

    if not storage_service.notebooklm_guideline_exists(transcription_id):
        raise HTTPException(
            status_code=404,
            detail="NotebookLM指南文件不存在。转录完成后会自动生成，请稍后再试。"
        )

    try:
        # Read guideline from storage
        content = storage_service.get_notebooklm_guideline(transcription_id)

        # Generate filename
        original_filename = Path(transcription.file_name).stem
        download_filename = f"{original_filename}-notebooklm.txt"

        from fastapi.responses import StreamingResponse
        from io import StringIO

        buffer = StringIO(content)

        return StreamingResponse(
            buffer,
            media_type="text/plain; charset=utf-8",
            headers={
                "Content-Disposition": f'attachment; filename="{download_filename}"'
            }
        )

    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="指南文件不存在")
    except Exception as e:
        logger.error(f"NotebookLM指南下载失败 {transcription_id}: {e}")
        raise HTTPException(status_code=500, detail=f"指南下载失败: {str(e)}")


# ==================== Chat Endpoints ====================

@router.get("/{transcription_id}/chat", response_model=ChatHistoryResponse)
async def get_chat_history(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    获取转录的聊天历史记录
    """
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    # Verify ownership
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="未找到转录")

    # Get chat history
    messages = db.query(ChatMessage).filter(
        ChatMessage.transcription_id == transcription_uuid
    ).order_by(ChatMessage.created_at).all()

    return ChatHistoryResponse(messages=messages)


@router.post("/{transcription_id}/chat", response_model=ChatMessageSchema)
async def send_chat_message(
    transcription_id: str,
    message: dict,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    发送聊天消息并获取AI回复
    """
    logger.info(f"[Chat] Received message for transcription {transcription_id}: {message}")

    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        logger.error(f"[Chat] Invalid UUID format: {transcription_id}")
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    # Verify ownership
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        logger.warning(f"[Chat] Transcription not found or access denied: {transcription_uuid}")
        raise HTTPException(status_code=404, detail="未找到转录")

    user_content = message.get("content")
    if not user_content:
        logger.warning("[Chat] Empty content received")
        raise HTTPException(status_code=400, detail="消息内容不能为空")

    logger.info(f"[Chat] Processing user message: {user_content[:100]}...")

    # Save user message
    user_message = ChatMessage(
        transcription_id=transcription_uuid,
        user_id=UUID(current_user.get("id")),
        role="user",
        content=user_content
    )
    db.add(user_message)
    db.commit()
    db.refresh(user_message)
    logger.info(f"[Chat] Saved user message: {user_message.id}")

    # Get chat history for context
    history = db.query(ChatMessage).filter(
        ChatMessage.transcription_id == transcription_uuid
    ).order_by(ChatMessage.created_at).all()

    chat_history = [
        {"role": msg.role, "content": msg.content}
        for msg in history[:-1]  # Exclude the message we just added
    ]
    logger.info(f"[Chat] Chat history length: {len(chat_history)}")

    # Get transcription text for context
    transcription_text = transcription.text
    logger.info(f"[Chat] Transcription text length: {len(transcription_text)} chars")

    # Call GLM API for response
    try:
        from app.core.glm import get_glm_client
        glm_client = get_glm_client()

        print(f"[Chat] Calling GLM API for transcription: {transcription_uuid}")
        logger.info("[Chat] Calling GLM API...")
        response = await glm_client.chat(
            question=user_content,
            transcription_context=transcription_text,
            chat_history=chat_history
        )

        assistant_content = response.get("response", "")
        logger.info(f"[Chat] GLM response received, length: {len(assistant_content)} chars")

    except Exception as e:
        logger.error(f"[Chat] GLM chat error: {e}", exc_info=True)
        assistant_content = "抱歉，AI回复失败，请稍后再试。"

    # Save assistant message
    assistant_message = ChatMessage(
        transcription_id=transcription_uuid,
        user_id=current_user.get("id"),
        role="assistant",
        content=assistant_content
    )
    db.add(assistant_message)
    db.commit()
    db.refresh(assistant_message)
    logger.info(f"[Chat] Saved assistant message: {assistant_message.id}")

    return assistant_message


@router.post("/{transcription_id}/chat/stream")
async def send_chat_message_stream(
    transcription_id: str,
    message: dict,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    发送聊天消息并获取AI回复（流式输出）

    使用Server-Sent Events (SSE)流式返回AI响应。
    """
    logger.info(f"[ChatStream] Received message for transcription {transcription_id}: {message}")

    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        logger.error(f"[ChatStream] Invalid UUID format: {transcription_id}")
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    # Verify ownership
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        logger.warning(f"[ChatStream] Transcription not found or access denied: {transcription_uuid}")
        raise HTTPException(status_code=404, detail="未找到转录")

    user_content = message.get("content")
    if not user_content:
        logger.warning("[ChatStream] Empty content received")
        raise HTTPException(status_code=400, detail="消息内容不能为空")

    logger.info(f"[ChatStream] Processing user message: {user_content[:100]}...")

    # Save user message immediately
    user_message = ChatMessage(
        transcription_id=transcription_uuid,
        user_id=current_user.get("id"),
        role="user",
        content=user_content
    )
    db.add(user_message)
    db.commit()
    db.refresh(user_message)
    logger.info(f"[ChatStream] Saved user message: {user_message.id}")

    # Get chat history for context
    history = db.query(ChatMessage).filter(
        ChatMessage.transcription_id == transcription_uuid
    ).order_by(ChatMessage.created_at).all()

    chat_history = [
        {"role": msg.role, "content": msg.content}
        for msg in history[:-1]  # Exclude the message we just added
    ]
    logger.info(f"[ChatStream] Chat history length: {len(chat_history)}")

    # Get transcription text for context
    transcription_text = transcription.text
    logger.info(f"[ChatStream] Transcription text length: {len(transcription_text)} chars")

    async def stream_generator():
        """Async generator wrapper for streaming response."""
        from app.core.glm import get_glm_client

        glm_client = get_glm_client()
        full_response = ""

        try:
            logger.info("[ChatStream] Starting stream from GLM API...")

            # chat_stream is a sync generator, iterate with regular for loop
            for chunk in glm_client.chat_stream(
                question=user_content,
                transcription_context=transcription_text,
                chat_history=chat_history
            ):
                # Yield each chunk immediately
                yield chunk

                # Parse the chunk to accumulate full response
                # Format: "data: {json}\n\n"
                if chunk.startswith("data: "):
                    json_str = chunk[6:].strip()  # Remove "data: " prefix
                    if json_str:
                        try:
                            import json
                            data = json.loads(json_str)
                            if "content" in data and not data.get("done"):
                                full_response += data["content"]
                        except json.JSONDecodeError:
                            pass

            # Save assistant message after stream completes
            logger.info(f"[ChatStream] Stream complete, saving assistant message (length: {len(full_response)})")
            assistant_message = ChatMessage(
                transcription_id=transcription_uuid,
                user_id=current_user.get("id"),
                role="assistant",
                content=full_response
            )
            db.add(assistant_message)
            db.commit()
            logger.info(f"[ChatStream] Saved assistant message: {assistant_message.id}")

        except Exception as e:
            logger.error(f"[ChatStream] Stream error: {e}", exc_info=True)
            # Send error through stream
            import json
            yield f"data: {json.dumps({'error': str(e), 'done': True})}\n\n"

            # Save error message
            assistant_message = ChatMessage(
                transcription_id=transcription_uuid,
                user_id=current_user.get("id"),
                role="assistant",
                content="抱歉，AI回复失败，请稍后再试。"
            )
            db.add(assistant_message)
            db.commit()

    return StreamingResponse(
        stream_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",  # Disable nginx buffering
        }
    )


# ==================== Share Link Endpoints ====================

def _generate_share_token() -> str:
    """Generate a secure URL-safe token for share links."""
    # Generate 16 random bytes and encode as URL-safe base64
    token_bytes = secrets.token_bytes(16)
    token = base64.urlsafe_b64encode(token_bytes).decode('utf-8').rstrip('=')
    return token


@router.post("/{transcription_id}/share", response_model=ShareLinkSchema)
async def create_share_link(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    创建转录的公开分享链接
    """
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    # Verify ownership
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="未找到转录")

    # Check if share link already exists
    existing_link = db.query(ShareLink).filter(
        ShareLink.transcription_id == transcription_uuid
    ).first()

    if existing_link:
        # Return existing link with URL
        from app.schemas.share import ShareLink as ShareLinkSchema
        schema_data = ShareLinkSchema.model_validate(existing_link).model_dump()
        schema_data['share_url'] = f"/shared/{existing_link.share_token}"
        return ShareLinkSchema(**schema_data)

    # Create new share link
    share_token = _generate_share_token()
    share_link = ShareLink(
        transcription_id=transcription_uuid,
        share_token=share_token
    )
    db.add(share_link)
    db.commit()
    db.refresh(share_link)

    # Return schema with share_url
    from app.schemas.share import ShareLink as ShareLinkSchema
    schema_data = ShareLinkSchema.model_validate(share_link).model_dump()
    schema_data['share_url'] = f"/shared/{share_token}"
    return ShareLinkSchema(**schema_data)


# ========================================
# Channel Assignment Endpoints
# ========================================

@router.post("/{transcription_id}/channels")
async def assign_transcription_to_channels(
    transcription_id: str,
    assignment: TranscriptionChannelAssignmentRequest,
    db: Session = Depends(get_db),
    current_db_user: User = Depends(get_current_db_user)
):
    """
    Assign transcription to channels (owner or admin only).

    Clears existing channel assignments and replaces with new ones.
    """
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid
    ).first()

    if not transcription:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Transcription not found"
        )

    # Only owner or admin can assign channels
    if transcription.user_id != current_db_user.id and not current_db_user.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to modify this transcription"
        )

    # Verify all channels exist
    channel_uuids = [UUID(str(cid)) for cid in assignment.channel_ids]
    channels = db.query(Channel).filter(Channel.id.in_(channel_uuids)).all()
    if len(channels) != len(channel_uuids):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="One or more channels not found"
        )

    # Remove existing assignments
    db.query(TranscriptionChannel).filter(
        TranscriptionChannel.transcription_id == transcription_uuid
    ).delete()

    # Add new assignments
    for channel_uuid in channel_uuids:
        ta = TranscriptionChannel(
            transcription_id=transcription_uuid,
            channel_id=channel_uuid,
            assigned_by=current_db_user.id
        )
        db.add(ta)

    db.commit()

    logger.info(
        f"Transcription {transcription_id} assigned to {len(channel_uuids)} channels "
        f"by user {current_db_user.email}"
    )
    return {
        "message": f"Assigned to {len(channel_uuids)} channels",
        "channel_ids": [str(cid) for cid in channel_uuids]
    }


@router.get("/{transcription_id}/channels", response_model=List[ChannelResponse])
async def get_transcription_channels(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_db_user: User = Depends(get_current_db_user)
):
    """
    Get channels assigned to a transcription.

    Returns channels the transcription is assigned to.
    Admin can see any transcription's channels, regular users only their own.
    """
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid
    ).first()

    if not transcription:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Transcription not found"
        )

    # Check access: owner or admin, or member of assigned channel
    if transcription.user_id != current_db_user.id and not current_db_user.is_admin:
        # Check if user is member of any channel this transcription is assigned to
        is_member = db.query(TranscriptionChannel).join(
            ChannelMembership,
            TranscriptionChannel.channel_id == ChannelMembership.channel_id
        ).filter(
            TranscriptionChannel.transcription_id == transcription_id,
            ChannelMembership.user_id == current_db_user.id
        ).first()
        if not is_member:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Not authorized to view this transcription"
            )

    # Get channel assignments
    assignments = db.query(TranscriptionChannel).filter(
        TranscriptionChannel.transcription_id == transcription_id
    ).all()

    channel_ids = [a.channel_id for a in assignments]
    channels = db.query(Channel).filter(Channel.id.in_(channel_ids)).all()

    # Convert SQLAlchemy models to Pydantic schemas
    return [ChannelResponse.model_validate(c) for c in channels]
