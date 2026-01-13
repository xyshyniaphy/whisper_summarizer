"""
Shared Transcription API Router

Public access endpoints for shared transcriptions (no authentication required).
"""

from fastapi import APIRouter, Depends, HTTPException, Query, BackgroundTasks, Request
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.orm import Session
from uuid import UUID
from pathlib import Path as FilePath
from urllib.parse import quote
from app.api.deps import get_db
from app.models.transcription import Transcription
from app.models.share_link import ShareLink
from app.models.chat_message import ChatMessage
from app.schemas.share import (
    SharedTranscriptionResponse,
)
from app.services.storage_service import get_storage_service
import logging
import aiofiles

logger = logging.getLogger(__name__)

router = APIRouter()


def _create_content_disposition(filename: str) -> dict:
    """
    Create Content-Disposition header with proper encoding for non-ASCII filenames.

    Uses RFC 2231 encoding for filenames with non-ASCII characters (e.g., Japanese).
    """
    if filename.isascii():
        return {"Content-Disposition": f'attachment; filename="{filename}"'}
    else:
        encoded_filename = quote(filename)
        return {"Content-Disposition": f'attachment; filename*=UTF-8\'\'{encoded_filename}'}


@router.get("/{share_token}", response_model=SharedTranscriptionResponse)
async def get_shared_transcription(
    share_token: str,
    db: Session = Depends(get_db)
):
    """
    访问公开分享的转录（无需认证）
    """
    # Find share link
    share_link = db.query(ShareLink).filter(
        ShareLink.share_token == share_token
    ).first()

    if not share_link:
        raise HTTPException(status_code=404, detail="分享链接不存在")

    # Check expiration
    if share_link.expires_at and share_link.expires_at < __import__('datetime').datetime.now(__import__('datetime').timezone.utc):
        raise HTTPException(status_code=410, detail="分享链接已过期")

    # Increment access count
    share_link.access_count += 1
    db.commit()

    # Get transcription
    transcription = db.query(Transcription).filter(
        Transcription.id == share_link.transcription_id
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="转录不存在")

    # Get summary
    summary = None
    if transcription.summaries and len(transcription.summaries) > 0:
        summary = transcription.summaries[0].summary_text

    # Get chat messages (sorted by created_at)
    chat_messages = db.query(ChatMessage).filter(
        ChatMessage.transcription_id == transcription.id
    ).order_by(ChatMessage.created_at).all()

    return SharedTranscriptionResponse(
        id=transcription.id,
        file_name=transcription.file_name,
        text=transcription.text,
        summary=summary,
        language=transcription.language,
        duration_seconds=transcription.duration_seconds,
        created_at=transcription.created_at,
        chat_messages=chat_messages
    )


@router.get("/{share_token}/download")
async def download_shared_transcription(
    share_token: str,
    format: str = Query("txt", enum=["txt", "srt"]),
    db: Session = Depends(get_db)
):
    """
    Download shared transcription file (public access, no authentication required).

    Supports txt and srt formats with proper Japanese filename encoding.
    """
    # Find share link
    share_link = db.query(ShareLink).filter(
        ShareLink.share_token == share_token
    ).first()

    if not share_link:
        raise HTTPException(status_code=404, detail="分享链接不存在")

    # Check expiration
    if share_link.expires_at and share_link.expires_at < __import__('datetime').datetime.now(__import__('datetime').timezone.utc):
        raise HTTPException(status_code=410, detail="分享链接已过期")

    # Get transcription
    transcription = db.query(Transcription).filter(
        Transcription.id == share_link.transcription_id
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="转录不存在")

    # Check if transcription has text
    if not transcription.text:
        raise HTTPException(status_code=400, detail="转录内容为空")

    # Get filename without extension
    original_filename = Path(transcription.file_name).stem

    # Generate content based on format
    if format == "srt":
        # Try to get segments for real timestamps
        storage_service = get_storage_service()
        if storage_service.segments_exist(str(transcription.id)):
            segments = storage_service.get_transcription_segments(str(transcription.id))
            if segments:
                # Use real timestamps from segments
                srt_lines = []
                for i, segment in enumerate(segments, 1):
                    srt_lines.append(f"{i}")
                    srt_lines.append(f"{segment['start']} --> {segment['end']}")
                    srt_lines.append(segment['text'])
                    srt_lines.append("")
                content = "\n".join(srt_lines)
                download_filename = f"{original_filename}.srt"
            else:
                # Fallback to fake timestamps
                from app.api.transcriptions import _format_fake_srt
                content = _format_fake_srt(transcription.text)
                download_filename = f"{original_filename}.srt"
        else:
            # No segments file - use fake timestamps
            from app.api.transcriptions import _format_fake_srt
            content = _format_fake_srt(transcription.text)
            download_filename = f"{original_filename}.srt"
    else:
        # TXT format - return the transcription text
        content = transcription.text
        download_filename = f"{original_filename}.txt"

    # Return file with proper encoding
    from io import StringIO

    buffer = StringIO(content)
    return StreamingResponse(
        buffer,
        media_type="text/plain; charset=utf-8",
        headers=_create_content_disposition(download_filename)
    )


@router.get("/{share_token}/segments")
async def get_shared_segments(
    share_token: str,
    db: Session = Depends(get_db)
):
    """
    Get transcription segments with timestamps for audio player navigation.

    Returns JSON array of segments with start, end, text fields.
    Returns empty array if segments file doesn't exist.
    """
    # Find share link
    share_link = db.query(ShareLink).filter(
        ShareLink.share_token == share_token
    ).first()

    if not share_link:
        raise HTTPException(status_code=404, detail="分享链接不存在")

    # Check expiration
    if share_link.expires_at and share_link.expires_at < __import__('datetime').datetime.now(__import__('datetime').timezone.utc):
        raise HTTPException(status_code=410, detail="分享链接已过期")

    # Get transcription
    transcription = db.query(Transcription).filter(
        Transcription.id == share_link.transcription_id
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="转录不存在")

    # Get segments from storage (returns empty list if not found)
    storage_service = get_storage_service()
    segments = storage_service.get_transcription_segments(str(transcription.id))

    return segments


async def _get_mime_type(file_path: str) -> str:
    """Get MIME type based on file extension."""
    ext = FilePath(file_path).suffix.lower()
    mime_types = {
        ".mp3": "audio/mpeg",
        ".m4a": "audio/mp4",
        ".wav": "audio/wav",
        ".ogg": "audio/ogg",
        ".webm": "audio/webm",
    }
    return mime_types.get(ext, "audio/mpeg")


@router.get("/{share_token}/audio")
async def get_shared_audio(
    share_token: str,
    request: Request,
    db: Session = Depends(get_db)
):
    """
    Stream original audio file for shared transcription.

    Supports HTTP Range requests for seeking in audio players.
    Public access (no authentication required).
    """
    # Find share link
    share_link = db.query(ShareLink).filter(
        ShareLink.share_token == share_token
    ).first()

    if not share_link:
        raise HTTPException(status_code=404, detail="分享链接不存在")

    # Check expiration
    if share_link.expires_at and share_link.expires_at < __import__('datetime').datetime.now(__import__('datetime').timezone.utc):
        raise HTTPException(status_code=410, detail="分享链接已过期")

    # Get transcription
    transcription = db.query(Transcription).filter(
        Transcription.id == share_link.transcription_id
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="转录不存在")

    # Check if file_path exists
    if not transcription.file_path:
        raise HTTPException(status_code=404, detail="音频文件不存在")

    file_path = FilePath(transcription.file_path)

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="音频文件未找到")

    # Get file size
    file_size = file_path.stat().st_size

    # Handle Range header
    range_header = request.headers.get("range")
    headers = {
        "content-type": await _get_mime_type(str(file_path)),
        "accept-ranges": "bytes",
    }

    if range_header:
        # Parse Range header (format: "bytes=start-end")
        try:
            range_match = range_header.replace("bytes=", "").strip()
            range_parts = range_match.split("-")
            start = int(range_parts[0]) if range_parts[0] else 0
            end = int(range_parts[1]) if range_parts[1] else file_size - 1

            # Validate range
            if start >= file_size or end >= file_size or start > end:
                raise HTTPException(
                    status_code=416,
                    detail="Invalid range",
                    headers={"content-range": f"bytes */{file_size}"}
                )

            # Read partial content
            chunk_size = end - start + 1
            async with aiofiles.open(file_path, "rb") as f:
                await f.seek(start)
                chunk = await f.read(chunk_size)

            headers["content-range"] = f"bytes {start}-{end}/{file_size}"
            headers["content-length"] = str(chunk_size)

            return StreamingResponse(
                iter([chunk]),
                status_code=206,
                headers=headers
            )
        except (ValueError, IndexError):
            raise HTTPException(status_code=400, detail="Invalid range header")

    # No Range header - return entire file
    headers["content-length"] = str(file_size)

    async def file_iterator():
        async with aiofiles.open(file_path, "rb") as f:
            while chunk := await f.read(64 * 1024):  # 64KB chunks
                yield chunk

    return StreamingResponse(
        file_iterator(),
        headers=headers
    )


@router.get("/{share_token}/download-docx")
async def download_shared_summary_docx(
    share_token: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    """
    Download shared transcription summary as DOCX file (public access, no authentication required).

    Generates a DOCX file from the AI summary using python-docx with proper Chinese font support.
    """
    # Find share link
    share_link = db.query(ShareLink).filter(
        ShareLink.share_token == share_token
    ).first()

    if not share_link:
        raise HTTPException(status_code=404, detail="分享链接不存在")

    # Check expiration
    if share_link.expires_at and share_link.expires_at < __import__('datetime').datetime.now(__import__('datetime').timezone.utc):
        raise HTTPException(status_code=410, detail="分享链接已过期")

    # Get transcription
    transcription = db.query(Transcription).filter(
        Transcription.id == share_link.transcription_id
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="转录不存在")

    # Get summary
    if not transcription.summaries or len(transcription.summaries) == 0:
        raise HTTPException(status_code=404, detail="未找到摘要数据")

    summary_text = transcription.summaries[0].summary_text
    if not summary_text:
        raise HTTPException(status_code=400, detail="摘要内容为空")

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import tempfile
        import re

        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        docx_path = tempfile.mktemp(suffix=".docx", dir=temp_dir)

        # 创建Document对象
        doc = Document()

        # 设置默认字体（支持中文）
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        doc.styles['Normal'].font.size = Pt(11)

        # 简单的Markdown解析
        lines = summary_text.split('\n')
        in_list = False

        for line in lines:
            line = line.rstrip()

            # 空行
            if not line:
                if in_list:
                    in_list = False
                doc.add_paragraph()
                continue

            # 标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 6:
                    heading_text = line.lstrip('#').strip()
                    heading = doc.add_heading(heading_text, level=min(level, 9))
                    # 设置中文字体
                    for run in heading.runs:
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    continue

            # 列表项
            if line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\. ', line):
                if not in_list:
                    in_list = True
                list_text = line.lstrip('-*+').lstrip('0123456789.')
                p = doc.add_paragraph(list_text, style='List Bullet')
                # 设置中文字体
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                continue

            # 代码块（简单处理）
            if line.startswith('```'):
                continue

            # 普通段落 - 处理内联格式
            in_list = False
            p = doc.add_paragraph()

            # 处理粗体和斜体
            parts = re.split(r'(\*\*.*?\*\*|_.*?_)', line)
            for part in parts:
                if not part:
                    continue

                run = p.add_run(part)

                # 设置中文字体
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

                if part.startswith('**') and part.endswith('**'):
                    run.bold = True
                    run.text = part[2:-2]
                elif part.startswith('_') and part.endswith('_'):
                    run.italic = True
                    run.text = part[1:-1]

        # 保存文档
        doc.save(docx_path)

        # 设置文件名
        original_filename = Path(transcription.file_name).stem
        download_filename = f"{original_filename}-摘要.docx"

        # 定义清理函数（在响应发送后执行）
        def cleanup_temp_files():
            import shutil
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
                logger.debug(f"Cleaned up temp directory: {temp_dir}")
            except Exception as e:
                logger.warning(f"Failed to cleanup temp directory: {e}")

        # 添加后台任务在响应完成后清理
        background_tasks.add_task(cleanup_temp_files)

        return FileResponse(
            path=docx_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=_create_content_disposition(download_filename)
        )

    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx未安装，请联系管理员"
        )
    except Exception as e:
        logger.error(f"DOCX生成失败 {transcription.id}: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX生成失败: {str(e)}")
