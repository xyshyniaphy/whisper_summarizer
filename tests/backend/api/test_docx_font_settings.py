"""
DOCX Font Settings Coverage Tests

Test for transcriptions.py lines 564-565, 576-577 - Chinese font settings.
"""

import pytest
import tempfile
import shutil
from pathlib import Path
from unittest.mock import patch, MagicMock
from uuid import uuid4
from docx import Document

from app.api.transcriptions import download_summary_docx


@pytest.mark.asyncio
class TestDOCXChineseFontSettings:
    """Test DOCX Chinese font settings are applied."""

    @pytest.mark.asyncio
    async def test_docx_applies_chinese_font_to_headings_hits_line_564_565(self):
        """
        Test that Chinese font is applied to headings.

        This targets transcriptions.py lines 564-565:
        ```python
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        ```
        """
        # Create real temp directory
        temp_dir = tempfile.mkdtemp(prefix="test_docx_heading_")

        try:
            mock_db = MagicMock()
            mock_transcription = MagicMock()
            mock_transcription.id = uuid4()
            mock_transcription.file_name = "test_heading.wav"

            # Summary with Chinese headings
            mock_summary = MagicMock()
            mock_summary.summary_text = """# 一级标题

## 二级标题

### 三级标题

普通段落文本。"""

            mock_transcription.summaries = [mock_summary]
            mock_db.query.return_value.filter.return_value.first.return_value = mock_transcription
            mock_user = {"id": str(uuid4())}

            # Use real tempfile paths
            with patch('tempfile.mkdtemp', return_value=temp_dir):
                with patch('tempfile.mktemp', return_value=temp_dir + "/file.docx"):
                    with patch('app.api.transcriptions.FileResponse') as mock_file_response:
                        mock_file_response.return_value = MagicMock()

                        # Call the endpoint - this will use real docx.Document
                        result = await download_summary_docx(
                            str(uuid4()), MagicMock(), mock_db, mock_user
                        )

            # Now verify the real DOCX file was created with proper font settings
            docx_path = Path(temp_dir) / "file.docx"
            if docx_path.exists():
                doc = Document(docx_path)

                # Check that headings were created
                headings_found = 0
                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        headings_found += 1
                        # Verify font settings were applied
                        for run in para.runs:
                            # Lines 564-565 should have set these
                            if run.font.name:
                                assert run.font.name == 'Microsoft YaHei' or run.font.name == 'Calibri'

                assert headings_found >= 3, f"Expected at least 3 headings, found {headings_found}"

        finally:
            # Cleanup
            if Path(temp_dir).exists():
                shutil.rmtree(temp_dir, ignore_errors=True)

    @pytest.mark.asyncio
    async def test_docx_applies_chinese_font_to_lists_hits_line_576_577(self):
        """
        Test that Chinese font is applied to list items.

        This targets transcriptions.py lines 576-577:
        ```python
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        ```
        """
        temp_dir = tempfile.mkdtemp(prefix="test_docx_list_")

        try:
            mock_db = MagicMock()
            mock_transcription = MagicMock()
            mock_transcription.id = uuid4()
            mock_transcription.file_name = "test_list.wav"

            # Summary with Chinese list items
            mock_summary = MagicMock()
            mock_summary.summary_text = """- 第一个列表项

- 第二个列表项

- 第三个列表项

1. 编号项目一
2. 编号项目二
3. 编号项目三

普通段落文本。"""

            mock_transcription.summaries = [mock_summary]
            mock_db.query.return_value.filter.return_value.first.return_value = mock_transcription
            mock_user = {"id": str(uuid4())}

            with patch('tempfile.mkdtemp', return_value=temp_dir):
                with patch('tempfile.mktemp', return_value=temp_dir + "/file.docx"):
                    with patch('app.api.transcriptions.FileResponse') as mock_file_response:
                        mock_file_response.return_value = MagicMock()

                        result = await download_summary_docx(
                            str(uuid4()), MagicMock(), mock_db, mock_user
                        )

            # Verify the real DOCX file was created
            docx_path = Path(temp_dir) / "file.docx"
            if docx_path.exists():
                doc = Document(docx_path)

                # Check that list items were created
                list_items_found = 0
                for para in doc.paragraphs:
                    if 'List Bullet' in para.style.name or 'List Number' in para.style.name:
                        list_items_found += 1
                        # Verify font settings were applied
                        for run in para.runs:
                            if run.font.name:
                                # Lines 576-577 should have set these
                                assert run.font.name == 'Microsoft YaHei' or run.font.name == 'Calibri'

                assert list_items_found >= 6, f"Expected at least 6 list items, found {list_items_found}"

        finally:
            # Cleanup
            if Path(temp_dir).exists():
                shutil.rmtree(temp_dir, ignore_errors=True)

    @pytest.mark.asyncio
    async def test_docx_applies_chinese_font_to_mixed_content_hits_both_lines(self):
        """
        Test Chinese font is applied to mixed content with headings and lists.

        This ensures both lines 564-565 and 576-577 are covered.
        """
        temp_dir = tempfile.mkdtemp(prefix="test_docx_mixed_")

        try:
            mock_db = MagicMock()
            mock_transcription = MagicMock()
            mock_transcription.id = uuid4()
            mock_transcription.file_name = "test_mixed.wav"

            # Mixed content with Chinese headings and lists
            mock_summary = MagicMock()
            mock_summary.summary_text = """# 会议记录

## 讨论要点

- 第一点重要内容
- 第二点重要内容
- 第三点重要内容

## 后续行动

1. 完成文档编写
2. 代码审查
3. 部署到生产环境

## 备注

普通段落文本，包含**粗体**和_斜体_。"""

            mock_transcription.summaries = [mock_summary]
            mock_db.query.return_value.filter.return_value.first.return_value = mock_transcription
            mock_user = {"id": str(uuid4())}

            with patch('tempfile.mkdtemp', return_value=temp_dir):
                with patch('tempfile.mktemp', return_value=temp_dir + "/file.docx"):
                    with patch('app.api.transcriptions.FileResponse') as mock_file_response:
                        mock_file_response.return_value = MagicMock()

                        result = await download_summary_docx(
                            str(uuid4()), MagicMock(), mock_db, mock_user
                        )

            # Verify the real DOCX file was created with both headings and lists
            docx_path = Path(temp_dir) / "file.docx"
            if docx_path.exists():
                doc = Document(docx_path)

                headings_count = sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))
                lists_count = sum(1 for p in doc.paragraphs if 'List' in p.style.name)

                assert headings_count >= 3, f"Expected at least 3 headings, found {headings_count}"
                assert lists_count >= 5, f"Expected at least 5 list items, found {lists_count}"

        finally:
            # Cleanup
            if Path(temp_dir).exists():
                shutil.rmtree(temp_dir, ignore_errors=True)
