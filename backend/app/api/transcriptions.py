from fastapi import APIRouter, Depends, HTTPException, Query, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from sqlalchemy.orm import Session
from typing import List, Optional
from pathlib import Path
from uuid import UUID
import secrets
import base64
from app.api.deps import get_db
from app.core.supabase import get_current_active_user
from app.models.transcription import Transcription
from app.models.summary import Summary
from app.models.chat_message import ChatMessage
from app.models.share_link import ShareLink
from app.schemas.transcription import Transcription as TranscriptionSchema
from app.schemas.summary import Summary as SummarySchema
from app.schemas.chat import ChatMessage as ChatMessageSchema, ChatHistoryResponse
from app.schemas.share import ShareLink as ShareLinkSchema
from app.services.pptx_service import get_pptx_service
from app.services.marp_service import get_marp_service

import logging
import asyncio
import time

logger = logging.getLogger(__name__)

router = APIRouter()


@router.get("", response_model=List[TranscriptionSchema])
async def list_transcriptions(
    limit: int = Query(10, ge=1, le=100),
    offset: int = Query(0, ge=0),
    stage: str = Query(None, description="Filter by stage: uploading, transcribing, summarizing, completed, failed"),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    文字起こしリストを取得 (現在のユーザーのもののみ)
    """
    query = db.query(Transcription).filter(Transcription.user_id == current_user.get("id"))

    if stage:
        query = query.filter(Transcription.stage == stage)
    
    transcriptions = query.order_by(Transcription.created_at.desc()).offset(offset).limit(limit).all()
    return transcriptions


@router.get("/{transcription_id}", response_model=TranscriptionSchema)
async def get_transcription(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    文字起こし詳細を取得
    """
    # Convert string ID to UUID for database query
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()
    
    if not transcription:
        raise HTTPException(status_code=404, detail="文字起こしが見つかりません")
    
    return transcription


@router.delete("/all", status_code=200)
async def delete_all_transcriptions(
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    すべての文字起こしを削除 (ファイルも含む)
    処理中の場合はキャンセルしてプロセスを終了してから削除します
    """
    from app.services.transcription_processor import (
        is_transcription_active,
        mark_transcription_cancelled,
        kill_transcription_processes,
        get_transcription_task_info
    )

    # Get all user's transcriptions
    transcriptions = db.query(Transcription).filter(
        Transcription.user_id == current_user.get("id")
    ).all()

    if not transcriptions:
        return {"deleted_count": 0, "message": "削除する項目がありません"}

    deleted_count = 0
    output_dir = Path("/app/data/output")

    for transcription in transcriptions:
        transcription_id = str(transcription.id)

        # Cancel active transcriptions
        if is_transcription_active(transcription_id):
            logger.info(f"[DELETE ALL] Cancelling active transcription: {transcription_id}")
            mark_transcription_cancelled(transcription_id)
            kill_transcription_processes(transcription_id)
            time.sleep(0.1)  # Brief pause between cancellations

        # Delete files
        try:
            import os
            from app.services.storage_service import get_storage_service

            # Delete from Supabase Storage
            if transcription.storage_path:
                try:
                    storage_service = get_storage_service()
                    storage_service.delete_transcription_text(str(transcription.id))
                    logger.info(f"[DELETE ALL] Deleted from Supabase Storage: {transcription.storage_path}")
                except Exception as e:
                    logger.warning(f"[DELETE ALL] Failed to delete from storage: {e}")

            if transcription.file_path and os.path.exists(transcription.file_path):
                os.remove(transcription.file_path)

            for ext in [".wav", ".txt", ".srt", ".vtt", ".json", ".pptx", ".md"]:
                output_file = output_dir / f"{transcription.id}{ext}"
                if output_file.exists():
                    output_file.unlink()

                converted_wav = output_dir / f"{transcription.id}_converted.wav"
                if converted_wav.exists():
                    converted_wav.unlink()
        except Exception as e:
            logger.error(f"[DELETE ALL] File deletion error for {transcription_id}: {e}")

        db.delete(transcription)
        deleted_count += 1

    db.commit()
    logger.info(f"[DELETE ALL] Deleted {deleted_count} transcriptions for user {current_user.get('id')}")

    return {
        "deleted_count": deleted_count,
        "message": f"{deleted_count}件の転写を削除しました"
    }


@router.delete("/{transcription_id}", status_code=204)
async def delete_transcription(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    文字起こしを削除 (ファイルも含む)
    処理中の場合はキャンセルしてプロセスを終了してから削除します
    """
    # Convert string ID to UUID for database query
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="文字起こしが見つかりません")

    # Check if transcription is currently processing and cancel if needed
    from app.services.transcription_processor import (
        is_transcription_active,
        mark_transcription_cancelled,
        kill_transcription_processes,
        get_transcription_task_info
    )

    if is_transcription_active(transcription_id):
        logger.info(f"[DELETE] Cancelling active transcription: {transcription_id}")

        # Get task info before cancelling
        task_info = get_transcription_task_info(transcription_id)
        if task_info:
            pids = task_info.get("pids", set())
            logger.info(f"[DELETE] Task info for {transcription_id}: Stage={task_info.get('stage')}, PIDs={list(pids)}")

        # Mark as cancelled (signals the transcription thread)
        was_marked = mark_transcription_cancelled(transcription_id)
        if was_marked:
            logger.info(f"[DELETE] Marked transcription {transcription_id} as cancelled")

        # Kill subprocesses (FFmpeg, whisper-cli)
        killed_count = kill_transcription_processes(transcription_id)
        logger.info(f"[DELETE] Killed {killed_count} subprocesses for transcription {transcription_id}")

        # Wait a moment for processes to terminate
        time.sleep(0.5)

    # ファイル削除
    try:
        import os
        from pathlib import Path
        from app.services.storage_service import get_storage_service

        # Delete from Supabase Storage
        if transcription.storage_path:
            try:
                storage_service = get_storage_service()
                storage_service.delete_transcription_text(str(transcription.id))
                logger.info(f"[DELETE] Deleted from Supabase Storage: {transcription.storage_path}")
            except Exception as e:
                logger.warning(f"[DELETE] Failed to delete from storage: {e}")

        # アップロードファイル削除
        if transcription.file_path and os.path.exists(transcription.file_path):
            os.remove(transcription.file_path)
            logger.info(f"[DELETE] Deleted upload file: {transcription.file_path}")

        # 出力ファイル削除 (wav, txt, srt, vtt, json, pptx, md)
        output_dir = Path("/app/data/output")
        for ext in [".wav", ".txt", ".srt", ".vtt", ".json", ".pptx", ".md"]:
            output_file = output_dir / f"{transcription.id}{ext}"
            if output_file.exists():
                output_file.unlink()
                logger.info(f"[DELETE] Deleted output file: {output_file}")

            # Also check for converted wav
            converted_wav = output_dir / f"{transcription.id}_converted.wav"
            if converted_wav.exists():
                converted_wav.unlink()
                logger.info(f"[DELETE] Deleted converted wav: {converted_wav}")

    except Exception as e:
        logger.error(f"ファイル削除エラー: {e}")
        # DB削除は続行する
        pass

    db.delete(transcription)
    db.commit()
    logger.info(f"[DELETE] Deleted transcription {transcription_id} from database")
    return None


@router.get("/{transcription_id}/download")
async def download_transcription(
  transcription_id: str,
  format: str = Query("txt", pattern="^(txt|srt|pptx)$"),
  db: Session = Depends(get_db),
  current_user: dict = Depends(get_current_active_user)
):
  """
  下载转录文件

  文件按需生成:
  - txt/srt: 从存储的转录文本生成
  - pptx: 检查现有文件

  Args:
    transcription_id: 转录ID
    format: 文件格式 (txt, srt 或 pptx)

  Returns:
    StreamingResponse: 下载文件
  """
  # 认证确认 - Convert string ID to UUID for database query
  try:
    transcription_uuid = UUID(transcription_id)
  except ValueError:
    raise HTTPException(status_code=422, detail="Invalid transcription ID format")
  transcription = db.query(Transcription).filter(
    Transcription.id == transcription_uuid,
    Transcription.user_id == current_user.get("id")
  ).first()

  if not transcription:
    raise HTTPException(status_code=404, detail="未找到转录")

  # PPTX 特殊处理 - 检查现有文件
  if format == "pptx":
    output_dir = Path("/app/data/output")
    file_path = output_dir / f"{transcription_id}.{format}"

    if not file_path.exists():
      raise HTTPException(status_code=404, detail="PPTX文件未找到，请先生成")

    original_filename = Path(transcription.file_name).stem
    download_filename = f"{original_filename}.{format}"

    return FileResponse(
      path=str(file_path),
      filename=download_filename,
      media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation"
    )

  # TXT/SRT - 从存储的转录文本按需生成
  if not transcription.text:
    raise HTTPException(status_code=400, detail="转录内容为空")

  content = transcription.text
  original_filename = Path(transcription.file_name).stem

  # SRT 格式 - 生成基本字幕格式
  if format == "srt":
    # 简单的 SRT 格式: 每行作为一个字幕段
    lines = content.strip().split("\n")
    srt_lines = []
    for i, line in enumerate(lines, 1):
      if line.strip():
        # 每行分配1秒时间 (简化版，因为没有时间戳信息)
        start_time = f"{i:02d}:00:00,000"
        end_time = f"{i:02d}:00:01,000"
        srt_lines.append(f"{i}\n{start_time} --> {end_time}\n{line}\n")
    content = "\n".join(srt_lines)
    download_filename = f"{original_filename}.srt"
  else:
    download_filename = f"{original_filename}.txt"

  # 使用 StreamingResponse 返回内存内容
  from fastapi.responses import StreamingResponse
  from io import StringIO

  buffer = StringIO(content)

  return StreamingResponse(
    buffer,
    media_type="text/plain; charset=utf-8",
    headers={
      "Content-Disposition": f'attachment; filename="{download_filename}"'
    }
  )


async def _generate_pptx_task(transcription_id: str, db: Session) -> None:
  """
  后台任务：使用Marp生成PPTX文件（AI驱动结构化）

  使用Gemini AI智能提取主题、生成目录、总结和后续安排。

  Args:
    transcription_id: 转录ID
    db: 数据库会话
  """
  # Set status to generating - Convert string ID to UUID for database query
  try:
    transcription_uuid = UUID(transcription_id)
  except ValueError:
    logger.error(f"Invalid transcription ID format: {transcription_id}")
    return
  transcription = db.query(Transcription).filter(
    Transcription.id == transcription_uuid
  ).first()

  if not transcription:
    logger.error(f"转录 {transcription_id} 未找到")
    return

  try:
    # Update status to generating
    transcription.pptx_status = "generating"
    transcription.pptx_error_message = None
    db.commit()

    if not transcription.text:
      raise ValueError("转录内容为空")

    # 获取摘要
    summary_text = None
    if transcription.summaries and len(transcription.summaries) > 0:
      summary_text = transcription.summaries[0].summary_text

    # 使用AI生成Marp Markdown结构
    marp_service = get_marp_service()
    markdown = await marp_service.generate_markdown(transcription, summary_text)

    # 保存Markdown文件
    md_path = marp_service.save_markdown(transcription_id, markdown)
    logger.info(f"Markdown生成成功: {md_path}")

    # 转换为PPTX
    pptx_path = marp_service.convert_to_pptx(md_path)
    logger.info(f"PPTX转换成功: {pptx_path}")

    # Update status to ready
    transcription.pptx_status = "ready"
    transcription.pptx_error_message = None
    db.commit()

    logger.info(f"PPTX生成成功 (AI结构化): {transcription_id}.pptx")

  except Exception as e:
    # Update status to error
    transcription.pptx_status = "error"
    transcription.pptx_error_message = str(e)
    db.commit()
    logger.error(f"PPTX生成失败 {transcription_id}: {e}")


@router.post("/{transcription_id}/generate-pptx")
async def generate_pptx(
  transcription_id: str,
  background_tasks: BackgroundTasks,
  db: Session = Depends(get_db),
  current_user: dict = Depends(get_current_active_user)
):
  """
  生成PowerPoint演示文稿 (使用Marp CLI)

  在后台生成包含转录内容和AI摘要的PPTX文件。

  Args:
    transcription_id: 转录ID

  Returns:
    JSONResponse: 生成状态
  """
  # 验证转录所有权 - Convert string ID to UUID for database query
  try:
    transcription_uuid = UUID(transcription_id)
  except ValueError:
    raise HTTPException(status_code=422, detail="Invalid transcription ID format")
  transcription = db.query(Transcription).filter(
    Transcription.id == transcription_uuid,
    Transcription.user_id == current_user.get("id")
  ).first()

  if not transcription:
    raise HTTPException(status_code=404, detail="未找到转录")

  if not transcription.text:
    raise HTTPException(status_code=400, detail="转录内容为空，无法生成PPT")

  # Check current status from database
  current_status = transcription.pptx_status or "not-started"

  # If already generating, return current status
  if current_status == "generating":
    return JSONResponse(
      status_code=202,
      content={
        "status": "generating",
        "message": "PPTX生成中..."
      }
    )

  # If already ready, verify file exists
  if current_status == "ready":
    marp_service = get_marp_service()
    if marp_service.pptx_exists(transcription_id):
      return JSONResponse(
        status_code=200,
        content={
          "status": "ready",
          "message": "PPTX文件已存在"
        }
      )
    # File missing, reset status
    transcription.pptx_status = "not-started"
    db.commit()

  # If previous error, reset to allow retry
  if current_status == "error":
    transcription.pptx_status = "not-started"
    transcription.pptx_error_message = None
    db.commit()

  # 添加后台任务
  background_tasks.add_task(_generate_pptx_task, transcription_id, db)

  return JSONResponse(
    status_code=202,
    content={
      "status": "generating",
      "message": "PPTX生成任务已启动"
    }
  )


@router.get("/{transcription_id}/pptx-status")
async def get_pptx_status(
  transcription_id: str,
  db: Session = Depends(get_db),
  current_user: dict = Depends(get_current_active_user)
):
  """
  检查PPTX文件生成状态

  从数据库字段读取状态，而不是检查文件是否存在。

  Args:
    transcription_id: 转录ID

  Returns:
    JSONResponse: PPTX状态
  """
  # 验证转录所有权 - Convert string ID to UUID for database query
  try:
    transcription_uuid = UUID(transcription_id)
  except ValueError:
    raise HTTPException(status_code=422, detail="Invalid transcription ID format")
  transcription = db.query(Transcription).filter(
    Transcription.id == transcription_uuid,
    Transcription.user_id == current_user.get("id")
  ).first()

  if not transcription:
    raise HTTPException(status_code=404, detail="未找到转录")

  # Read status from database field
  status = transcription.pptx_status or "not-started"

  # Also verify file exists for 'ready' status
  marp_service = get_marp_service()
  exists = marp_service.pptx_exists(transcription_id)

  # If status says ready but file doesn't exist, reset to not-started
  if status == "ready" and not exists:
    status = "not-started"

  return JSONResponse(
    content={
      "status": status,
      "exists": exists
    }
  )


@router.get("/{transcription_id}/markdown")
async def get_markdown(
  transcription_id: str,
  db: Session = Depends(get_db),
  current_user: dict = Depends(get_current_active_user)
):
  """
  获取Marp Markdown内容

  使用AI生成结构化的Markdown内容（包含主题、目录、总结、后续安排）。
  如果markdown文件不存在，则重新生成。

  Args:
    transcription_id: 转录ID

  Returns:
    JSONResponse: Markdown内容
  """
  # 验证转录所有权 - Convert string ID to UUID for database query
  try:
    transcription_uuid = UUID(transcription_id)
  except ValueError:
    raise HTTPException(status_code=422, detail="Invalid transcription ID format")
  transcription = db.query(Transcription).filter(
    Transcription.id == transcription_uuid,
    Transcription.user_id == current_user.get("id")
  ).first()

  if not transcription:
    raise HTTPException(status_code=404, detail="未找到转录")

  if not transcription.text:
    raise HTTPException(status_code=400, detail="转录内容为空")

  marp_service = get_marp_service()

  # Check if markdown exists
  if marp_service.markdown_exists(transcription_id):
    md_path = marp_service.get_markdown_path(transcription_id)
    markdown = md_path.read_text(encoding="utf-8")
    return JSONResponse(
      content={
        "markdown": markdown,
        "cached": True
      }
    )

  # Generate new markdown
  try:
    # 获取摘要
    summary_text = None
    if transcription.summaries and len(transcription.summaries) > 0:
      summary_text = transcription.summaries[0].summary_text

    markdown = await marp_service.generate_markdown(transcription, summary_text)

    # Save markdown for future use
    marp_service.save_markdown(transcription_id, markdown)

    return JSONResponse(
      content={
        "markdown": markdown,
        "cached": False
      }
    )

  except Exception as e:
    logger.error(f"Markdown生成失败 {transcription_id}: {e}")
    raise HTTPException(status_code=500, detail=f"Markdown生成失败: {str(e)}")


@router.get("/{transcription_id}/download-markdown")
async def download_markdown(
  transcription_id: str,
  db: Session = Depends(get_db),
  current_user: dict = Depends(get_current_active_user)
):
  """
  下载Marp Markdown文件

  Args:
    transcription_id: 转录ID

  Returns:
    FileResponse: Markdown文件下载
  """
  # 验证转录所有权 - Convert string ID to UUID for database query
  try:
    transcription_uuid = UUID(transcription_id)
  except ValueError:
    raise HTTPException(status_code=422, detail="Invalid transcription ID format")
  transcription = db.query(Transcription).filter(
    Transcription.id == transcription_uuid,
    Transcription.user_id == current_user.get("id")
  ).first()

  if not transcription:
    raise HTTPException(status_code=404, detail="未找到转录")

  marp_service = get_marp_service()
  md_path = marp_service.get_markdown_path(transcription_id)

  # Generate if doesn't exist
  if not md_path.exists():
    try:
      # 获取摘要
      summary_text = None
      if transcription.summaries and len(transcription.summaries) > 0:
        summary_text = transcription.summaries[0].summary_text

      markdown = await marp_service.generate_markdown(transcription, summary_text)
      marp_service.save_markdown(transcription_id, markdown)
    except Exception as e:
      raise HTTPException(status_code=500, detail=f"Markdown生成失败: {str(e)}")

  # 设置文件名
  original_filename = Path(transcription.file_name).stem
  download_filename = f"{original_filename}-marp.md"

  return FileResponse(
    path=md_path,
    filename=download_filename,
    media_type="text/markdown"
  )


@router.get("/{transcription_id}/download-docx")
async def download_summary_docx(
    transcription_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    下载AI摘要的DOCX文件

    使用python-docx将摘要Markdown转换为DOCX格式，支持中文。

    Args:
        transcription_id: 转录ID

    Returns:
        FileResponse: DOCX文件下载
    """
    # 验证转录所有权 - Convert string ID to UUID for database query
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="未找到转录")

    # 获取摘要
    if not transcription.summaries or len(transcription.summaries) == 0:
        raise HTTPException(status_code=404, detail="未找到摘要数据")

    summary_text = transcription.summaries[0].summary_text
    if not summary_text:
        raise HTTPException(status_code=400, detail="摘要内容为空")

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import tempfile
        import re

        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        docx_path = tempfile.mktemp(suffix=".docx", dir=temp_dir)

        # 创建Document对象
        doc = Document()

        # 设置默认字体（支持中文）
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        doc.styles['Normal'].font.size = Pt(11)

        # 简单的Markdown解析
        lines = summary_text.split('\n')
        in_list = False

        for line in lines:
            line = line.rstrip()

            # 空行
            if not line:
                if in_list:
                    in_list = False
                doc.add_paragraph()
                continue

            # 标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 6:
                    heading_text = line.lstrip('#').strip()
                    heading = doc.add_heading(heading_text, level=min(level, 9))
                    # 设置中文字体
                    for run in heading.runs:
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    continue

            # 列表项
            if line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\. ', line):
                if not in_list:
                    in_list = True
                list_text = line.lstrip('-*+').lstrip('0123456789.')
                p = doc.add_paragraph(list_text, style='List Bullet')
                # 设置中文字体
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                continue

            # 代码块（简单处理）
            if line.startswith('```'):
                continue

            # 普通段落 - 处理内联格式
            in_list = False
            p = doc.add_paragraph()

            # 处理粗体和斜体
            parts = re.split(r'(\*\*.*?\*\*|_.*?_)', line)
            for part in parts:
                if not part:
                    continue

                run = p.add_run(part)

                # 设置中文字体
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

                if part.startswith('**') and part.endswith('**'):
                    run.bold = True
                    run.text = part[2:-2]
                elif part.startswith('_') and part.endswith('_'):
                    run.italic = True
                    run.text = part[1:-1]

        # 保存文档
        doc.save(docx_path)

        # 设置文件名
        original_filename = Path(transcription.file_name).stem
        download_filename = f"{original_filename}-摘要.docx"

        # 定义清理函数（在响应发送后执行）
        def cleanup_temp_files():
            import shutil
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
                logger.debug(f"Cleaned up temp directory: {temp_dir}")
            except Exception as e:
                logger.warning(f"Failed to cleanup temp directory: {e}")

        # 添加后台任务在响应完成后清理
        background_tasks.add_task(cleanup_temp_files)

        return FileResponse(
            path=docx_path,
            filename=download_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx未安装，请联系管理员"
        )
    except Exception as e:
        logger.error(f"DOCX生成失败 {transcription_id}: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX生成失败: {str(e)}")


# ==================== Chat Endpoints ====================

@router.get("/{transcription_id}/chat", response_model=ChatHistoryResponse)
async def get_chat_history(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    获取转录的聊天历史记录
    """
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    # Verify ownership
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="未找到转录")

    # Get chat history
    messages = db.query(ChatMessage).filter(
        ChatMessage.transcription_id == transcription_uuid
    ).order_by(ChatMessage.created_at).all()

    return ChatHistoryResponse(messages=messages)


@router.post("/{transcription_id}/chat", response_model=ChatMessageSchema)
async def send_chat_message(
    transcription_id: str,
    message: dict,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    发送聊天消息并获取AI回复
    """
    logger.info(f"[Chat] Received message for transcription {transcription_id}: {message}")

    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        logger.error(f"[Chat] Invalid UUID format: {transcription_id}")
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    # Verify ownership
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        logger.warning(f"[Chat] Transcription not found or access denied: {transcription_uuid}")
        raise HTTPException(status_code=404, detail="未找到转录")

    user_content = message.get("content")
    if not user_content:
        logger.warning("[Chat] Empty content received")
        raise HTTPException(status_code=400, detail="消息内容不能为空")

    logger.info(f"[Chat] Processing user message: {user_content[:100]}...")

    # Save user message
    user_message = ChatMessage(
        transcription_id=transcription_uuid,
        user_id=current_user.get("id"),
        role="user",
        content=user_content
    )
    db.add(user_message)
    db.commit()
    db.refresh(user_message)
    logger.info(f"[Chat] Saved user message: {user_message.id}")

    # Get chat history for context
    history = db.query(ChatMessage).filter(
        ChatMessage.transcription_id == transcription_uuid
    ).order_by(ChatMessage.created_at).all()

    chat_history = [
        {"role": msg.role, "content": msg.content}
        for msg in history[:-1]  # Exclude the message we just added
    ]
    logger.info(f"[Chat] Chat history length: {len(chat_history)}")

    # Get transcription text for context
    transcription_text = transcription.text
    logger.info(f"[Chat] Transcription text length: {len(transcription_text)} chars")

    # Call Gemini API for response
    try:
        from app.core.gemini import get_gemini_client
        gemini_client = get_gemini_client()

        print(f"[Chat] Calling Gemini API for transcription: {transcription_uuid}")
        logger.info("[Chat] Calling Gemini API...")
        response = await gemini_client.chat(
            question=user_content,
            transcription_context=transcription_text,
            chat_history=chat_history
        )

        assistant_content = response.get("response", "")
        logger.info(f"[Chat] Gemini response received, length: {len(assistant_content)} chars")

    except Exception as e:
        logger.error(f"[Chat] Gemini chat error: {e}", exc_info=True)
        assistant_content = "抱歉，AI回复失败，请稍后再试。"

    # Save assistant message
    assistant_message = ChatMessage(
        transcription_id=transcription_uuid,
        user_id=current_user.get("id"),
        role="assistant",
        content=assistant_content
    )
    db.add(assistant_message)
    db.commit()
    db.refresh(assistant_message)
    logger.info(f"[Chat] Saved assistant message: {assistant_message.id}")

    return assistant_message


# ==================== Share Link Endpoints ====================

def _generate_share_token() -> str:
    """Generate a secure URL-safe token for share links."""
    # Generate 16 random bytes and encode as URL-safe base64
    token_bytes = secrets.token_bytes(16)
    token = base64.urlsafe_b64encode(token_bytes).decode('utf-8').rstrip('=')
    return token


@router.post("/{transcription_id}/share", response_model=ShareLinkSchema)
async def create_share_link(
    transcription_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_active_user)
):
    """
    创建转录的公开分享链接
    """
    try:
        transcription_uuid = UUID(transcription_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid transcription ID format")

    # Verify ownership
    transcription = db.query(Transcription).filter(
        Transcription.id == transcription_uuid,
        Transcription.user_id == current_user.get("id")
    ).first()

    if not transcription:
        raise HTTPException(status_code=404, detail="未找到转录")

    # Check if share link already exists
    existing_link = db.query(ShareLink).filter(
        ShareLink.transcription_id == transcription_uuid
    ).first()

    if existing_link:
        # Return existing link with updated URL
        existing_link.share_url = f"/shared/{existing_link.share_token}"
        return existing_link

    # Create new share link
    share_token = _generate_share_token()
    share_link = ShareLink(
        transcription_id=transcription_uuid,
        share_token=share_token
    )
    db.add(share_link)
    db.commit()
    db.refresh(share_link)

    # Add the share_url
    share_link.share_url = f"/shared/{share_token}"

    return share_link
